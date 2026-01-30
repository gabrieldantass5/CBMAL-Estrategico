#!/usr/bin/env python3
import sys
import os
import subprocess
import re
import markdown
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Caminho absoluto do Pandoc
PANDOC_PATH = r"C:\Users\D_A_N\AppData\Local\Microsoft\WinGet\Packages\JohnMacFarlane.Pandoc_Microsoft.Winget.Source_8wekyb3d8bbwe\pandoc-3.8.3\pandoc.exe"

def set_table_border(table):
    """Adiciona bordas pretas simples à tabela."""
    tbl = table._tbl
    for cell in tbl.xpath('.//w:tc'):
        tcPr = cell.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                              '<w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                              '<w:left w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                              '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                              '<w:right w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                              '</w:tcBorders>')
        tcPr.append(tcBorders)

def convert_to_docx(md_path):
    """Converte Markdown para DOCX no padrão ABNT/CBMAL com foco em tabelas."""
    if not os.path.exists(md_path): return
    docx_path = md_path.replace('.md', '.docx')
    
    # Comando Pandoc simplificado
    cmd = [PANDOC_PATH, md_path, "-o", docx_path, "--standalone"]
    
    try:
        subprocess.run(cmd, check=True)
        doc = Document(docx_path)
        
        # Margens ABNT
        for section in doc.sections:
            section.top_margin = Cm(3.0)
            section.left_margin = Cm(3.0)
            section.bottom_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

        # Formatação de Parágrafos
        for para in doc.paragraphs:
            if not para.text.strip(): continue
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(12)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            para.paragraph_format.first_line_indent = Cm(1.25)
            para.paragraph_format.space_after = Pt(6)

        # Formatação de Tabelas (O PONTO CRÍTICO)
        for table in doc.tables:
            set_table_border(table)
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11) # Tabelas costumam usar fonte levemente menor
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1.0

        doc.save(docx_path)
        print(f"✅ DOCX (Premium) gerado: {docx_path}")
        print(f"💡 Dica: Agora é só abrir no Google Docs. As tabelas estarão com bordas e fontes corrigidas.")
    except Exception as e:
        print(f"❌ Erro no DOCX: {e}")

def convert_to_sei(md_path):
    """Converte Markdown para HTML (Padrão SEI)."""
    if not os.path.exists(md_path): return
    html_path = md_path.replace('.md', '.html')
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            text = f.read()
            
        # Mermaid protection
        text = re.sub(r'```mermaid\s+(.*?)\s+```', '<div style="border: 2px dashed #cc0000; padding: 10px;">⚠️ DIAGRAMA MERMAID</div>', text, flags=re.DOTALL)
        
        html_body = markdown.markdown(text, extensions=['extra', 'nl2br'])
        # Inserindo CSS de tabela básico para o SEI
        table_style = "<style>table { border-collapse: collapse; width: 100%; } th, td { border: 1px solid black; padding: 8px; text-align: left; } th { background-color: #f2f2f2; }</style>"
        html_content = f"<html><head>{table_style}</head><body style='font-family: Arial; font-size: 12pt; text-align: justify;'>{html_body}</body></html>"
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
        print(f"✅ HTML (SEI) gerado: {html_path}")
    except Exception as e:
        print(f"❌ Erro no HTML: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Uso: python scripts/doc_converter.py <formato: docx|sei> <arquivo.md>")
    else:
        fmt, path = sys.argv[1], sys.argv[2]
        if fmt == "docx": convert_to_docx(path)
        elif fmt == "sei": convert_to_sei(path)
